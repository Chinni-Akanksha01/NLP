#!/usr/bin/env python
# coding: utf-8

# In[2]:


pip install PyPDF2


# In[4]:


import PyPDF2
from PyPDF2 import PdfFileReader


# In[5]:


PyPDF2.__version__


# In[8]:


pdf=open(r"C:\Users\chinn\Downloads\trainer.pdf","rb")
pdf_reader=PyPDF2.PdfReader(pdf)
print("Number of pages:",len(pdf_reader.pages))
page=pdf_reader.pages[1]
print(page.extract_text())
pdf.close()


# In[18]:


import PyPDF2,urllib,nltk
from io import BytesIO
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords


# In[19]:


wFile=urllib.request.urlopen('http://www.udri.org/pdf/02%20working%20paper%201.pdf')
pdfreader=PyPDF2.PdfReader(BytesIO(wFile.read()))


# In[17]:


pageObj=pdfreader.pages[2]
page2=pageObj.extract_text()
punctuations=['(',')',';',':','[',']',',','...','.']
tokens=word_tokenize(page2)
stop_words=stopwords.words('english')
keywords=[word for word in tokens if not word in stop_words and not word in punctuations]


# In[20]:


keywords


# In[21]:


name_list=list()
check=['Mr.','Mrs.','Ms.'] for idx,token in enumerate(tokens):
    if token.startswith(tuple(check)) and idx < (len(tokens)-1):
        name=token+tokens[idx+1]+''+tokens[idx+2]
        name_list.append(name)
        print(name_list)


# In[22]:


get_ipython().system('pip install python-docx')


# In[23]:


import docx


# In[26]:


doc=open(r"C:\Users\chinn\OneDrive\Documents\AT_16.docx","rb")
document=docx.Document(doc)


# In[27]:


docu=""
for para in document.paragraphs:
    docu+=para.text
print(docu)


# In[30]:


for i in range(len(document.paragraphs)):
    print("The content of the paragraph "+str(i)+" is : "+document.paragraphs[i].text+"\n")


# In[31]:


get_ipython().system('pip install bs4')


# In[32]:


import urllib.request as urllib2
from bs4 import BeautifulSoup


# In[33]:


response=urllib2.urlopen('https://en.wikipedia.org/wiki/Natural_language_processing')
html_doc=response.read()


# In[34]:


soup=BeautifulSoup(html_doc,'html.parser')
strhtm=soup.prettify()
print(strhtm[:5000])


# In[ ]:




